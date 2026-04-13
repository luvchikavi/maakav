"""Shared RTL utilities for Word report generation."""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def setup_rtl_document(doc: Document):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "David"
    font.size = Pt(11)
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)


def add_rtl_paragraph(doc, text, bold=False, size=11, alignment=None):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = alignment or WD_ALIGN_PARAGRAPH.RIGHT
    pPr = p._element.get_or_add_pPr()
    pPr.append(pPr.makeelement(qn("w:bidi"), {}))
    run = p.add_run(text)
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    run.font.name = "David"
    rPr = run._element.get_or_add_rPr()
    rPr.append(rPr.makeelement(qn("w:rtl"), {}))
    return p


def add_rtl_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = h._element.get_or_add_pPr()
    pPr.append(pPr.makeelement(qn("w:bidi"), {}))
    for run in h.runs:
        run.font.name = "David"
        rPr = run._element.get_or_add_rPr()
        rPr.append(rPr.makeelement(qn("w:rtl"), {}))


def create_rtl_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pPr = p._element.get_or_add_pPr()
            pPr.append(pPr.makeelement(qn("w:bidi"), {}))
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = "David"
    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(value)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pPr = p._element.get_or_add_pPr()
                pPr.append(pPr.makeelement(qn("w:bidi"), {}))
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "David"
    return table


def format_currency(value):
    if value is None:
        return "-"
    try:
        return f"{float(value):,.0f} ₪"
    except (ValueError, TypeError):
        return str(value)


def format_percent(value):
    if value is None:
        return "-"
    try:
        return f"{float(value):.1f}%"
    except (ValueError, TypeError):
        return str(value)
