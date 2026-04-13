"""
Tracking Report Generator - produces Word (docx) document
matching the Israeli banking standard monitoring report format.

RTL Hebrew document with chapters 4-11.
"""

import io
from datetime import date
from decimal import Decimal
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from .chapters.ch4_budget import add_chapter_4
from .chapters.ch5_construction import add_chapter_5
from .chapters.ch7_sales import add_chapter_7
from .chapters.ch8_vat import add_chapter_8
from .chapters.ch9_equity import add_chapter_9
from .chapters.ch10_profitability import add_chapter_10
from .chapters.ch11_sources_uses import add_chapter_11


def setup_rtl_document(doc: Document):
    """Configure the document for RTL Hebrew text."""
    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "David"
    font.size = Pt(11)

    # Set RTL for all paragraphs
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)


def add_rtl_paragraph(doc: Document, text: str, bold: bool = False, size: int = 11, alignment=None) -> None:
    """Add a right-to-left paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = alignment or WD_ALIGN_PARAGRAPH.RIGHT
    # Set RTL
    pPr = p._element.get_or_add_pPr()
    bidi = pPr.makeelement(qn("w:bidi"), {})
    pPr.append(bidi)

    run = p.add_run(text)
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    run.font.name = "David"
    # Set RTL on run
    rPr = run._element.get_or_add_rPr()
    rtl = rPr.makeelement(qn("w:rtl"), {})
    rPr.append(rtl)

    return p


def add_rtl_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a right-to-left heading."""
    h = doc.add_heading(text, level=level)
    h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = h._element.get_or_add_pPr()
    bidi = pPr.makeelement(qn("w:bidi"), {})
    pPr.append(bidi)

    for run in h.runs:
        run.font.name = "David"
        rPr = run._element.get_or_add_rPr()
        rtl = rPr.makeelement(qn("w:rtl"), {})
        rPr.append(rtl)


def create_rtl_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    """Create a right-to-left table with headers and data."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Headers
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pPr = paragraph._element.get_or_add_pPr()
            bidi = pPr.makeelement(qn("w:bidi"), {})
            pPr.append(bidi)
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = "David"

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pPr = paragraph._element.get_or_add_pPr()
                bidi = pPr.makeelement(qn("w:bidi"), {})
                pPr.append(bidi)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = "David"

    return table


def format_currency(value) -> str:
    """Format a number as Israeli currency."""
    if value is None:
        return "-"
    try:
        v = float(value)
        return f"{v:,.0f} ₪"
    except (ValueError, TypeError):
        return str(value)


def format_percent(value) -> str:
    """Format a number as percentage."""
    if value is None:
        return "-"
    try:
        return f"{float(value):.1f}%"
    except (ValueError, TypeError):
        return str(value)


def generate_tracking_report(
    project_data: dict,
    report_data: dict,
    calc_results: dict,
) -> io.BytesIO:
    """
    Generate the full tracking report as a Word document.

    Args:
        project_data: Project details (name, address, bank, etc.)
        report_data: Report metadata (number, month, index)
        calc_results: Results from the calculation engine

    Returns:
        BytesIO containing the .docx file
    """
    doc = Document()
    setup_rtl_document(doc)

    # Cover page
    doc.add_paragraph()  # spacing
    add_rtl_heading(doc, f'דו"ח מעקב מס\' {report_data["report_number"]}', level=1)
    add_rtl_paragraph(doc, f'ליום {report_data["report_month"]}', size=14)
    add_rtl_paragraph(doc, f'פרויקט: {project_data["project_name"]}', size=14, bold=True)
    if project_data.get("address"):
        add_rtl_paragraph(doc, f'{project_data["address"]}, {project_data.get("city", "")}', size=12)
    if project_data.get("developer_name"):
        add_rtl_paragraph(doc, f'היזם: {project_data["developer_name"]}', size=12)

    doc.add_page_break()

    # Chapter 4: Budget
    if "budget_tracking" in calc_results:
        add_chapter_4(doc, calc_results["budget_tracking"], report_data)

    doc.add_page_break()

    # Chapter 5: Construction
    if "construction" in calc_results:
        add_chapter_5(doc, calc_results["construction"], calc_results.get("budget_tracking"))

    doc.add_page_break()

    # Chapter 7: Sales
    if "sales" in calc_results:
        add_chapter_7(doc, calc_results["sales"])

    doc.add_page_break()

    # Chapter 8: VAT
    if "vat" in calc_results:
        add_chapter_8(doc, calc_results["vat"])

    # Chapter 9: Equity
    if "equity" in calc_results:
        add_chapter_9(doc, calc_results["equity"])

    doc.add_page_break()

    # Chapter 10: Profitability
    if "profitability" in calc_results:
        add_chapter_10(doc, calc_results["profitability"])

    # Chapter 11: Sources & Uses
    if "sources_uses" in calc_results:
        add_chapter_11(doc, calc_results["sources_uses"])

    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
